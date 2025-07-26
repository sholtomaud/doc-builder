# ./Iterate2.py (or your chosen filename)
import os
import subprocess
import time  # Add time module for delays

from docx import Document
from docx.shared import Inches, Mm
from PIL import Image  # Pillow library

# --- Path Configuration ---
script_dir = os.path.dirname(os.path.abspath(__file__))
output_dir_name = "output_files"
abs_output_dir = os.path.join(script_dir, output_dir_name)

if not os.path.exists(abs_output_dir):
    os.makedirs(abs_output_dir)
    print(f"Created output directory: {abs_output_dir}")
else:
    print(f"Output directory already exists: {abs_output_dir}")


# Helper to run rhinocode command
def run_rhinocode_command(command_str):
    rhino_executable_path = "/Applications/RhinoWIP.app/Contents/Resources/bin/rhinocode"
    if not os.path.exists(rhino_executable_path):
        rhino_executable_path = "/Applications/Rhino 8.app/Contents/Resources/bin/rhinocode"
        if not os.path.exists(rhino_executable_path):
            print("ERROR: Rhino executable not found at /Applications/RhinoWIP.app/... or /Applications/Rhino 8.app/...")
            print("Please verify the path to 'rhinocode'.")
            return None

    cmd = [rhino_executable_path, "command", command_str]

    full_command_for_print = f"Executing: {rhino_executable_path} command \"{command_str}\""
    print(full_command_for_print)

    process = subprocess.run(cmd, capture_output=True, text=True, check=False)

    stdout_lower = process.stdout.lower()
    stderr_lower = process.stderr.lower()
    problem_keywords = ["unknown command", "error:", "failed", "unable to", "exception:"]
    stdout_has_problem = any(keyword in stdout_lower for keyword in problem_keywords)
    stderr_has_problem = any(keyword in stderr_lower for keyword in problem_keywords)

    print("--- Rhino Command Output ---")
    print(f"Command Sent: \"{command_str}\"")
    print(f"Return Code: {process.returncode}")
    if process.stdout.strip():
        print(f"Stdout:\n{process.stdout.strip()}")
    else:
        print("Stdout: (empty)")
    if process.stderr.strip():
        print(f"Stderr:\n{process.stderr.strip()}")
    else:
        print("Stderr: (empty)")

    if process.returncode != 0 or stdout_has_problem or stderr_has_problem:
        print("Status: Potential Issue Detected in Rhino execution.")
    else:
        print("Status: Rhino command likely succeeded.")
    print("----------------------------")

    return process

def set_landscape_a3(doc):
    if not doc.sections:
        doc.add_section()
    section = doc.sections[0]
    try:
        from docx.enum.section import WD_ORIENT
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(420)
        section.page_height = Mm(297)
    except ImportError:
        print("WD_ORIENT not found, setting page dimensions for landscape manually.")
        section.page_height = Mm(297) # A3 landscape height
        section.page_width = Mm(420)  # A3 landscape width

def create_docx_from_image(image_path, docx_path, model_value):
    if not os.path.exists(image_path):
        print(f"Error in create_docx_from_image: Image not found at {image_path}. Cannot create DOCX.")
        return

    print(f"Creating DOCX: {docx_path} from image: {image_path}")
    doc = Document()
    set_landscape_a3(doc)
    doc.add_heading(f'Rhino Model View - Value {model_value}', level=1)
    try:
        with Image.open(image_path) as img:
            width_px, height_px = img.size
            # If ViewCaptureToFile doesn't embed DPI, Pillow might use a default like 72 or 96.
            # Forcing an interpretation for DOCX layout:
            assumed_dpi_for_layout = 150 # Higher DPI means smaller size in inches for same pixels

            img_width_inches_at_assumed_dpi = width_px / assumed_dpi_for_layout
            # img_height_inches_at_assumed_dpi = height_px / assumed_dpi_for_layout # Not directly used if only width is set

        # A3 Landscape: page width Mm(420), page height Mm(297)
        # Margins: typical 1 inch (25.4mm) each side. Or 0.5 inch for more space.
        page_margin_inches = 0.75 # Use Inches(0.75).inches if you prefer

        # Available width on A3 landscape page
        a3_width_inches = Mm(420).inches
        available_doc_width_inches = a3_width_inches - (2 * page_margin_inches)

        print(f"Image pixel dimensions: {width_px}x{height_px} px")
        print(f"Image size at {assumed_dpi_for_layout} DPI for layout: {img_width_inches_at_assumed_dpi:.2f} inches wide")
        print(f"DOCX available width (A3 landscape with {page_margin_inches}\" margins): {available_doc_width_inches:.2f} inches")

        # Determine the width to use for the image in the document
        # We want to fit it within the available_doc_width_inches
        # If the image (at assumed_dpi_for_layout) is already narrower than available, use its size.
        # Otherwise, shrink it to fit the available width.
        display_width_inches = min(img_width_inches_at_assumed_dpi, available_doc_width_inches)

        # Sanity check: if display_width_inches becomes too small (e.g. very high DPI image)
        # ensure it's at least a certain minimum or rethink. For now, this logic is usually fine.
        # If original pixel width is small, e.g. 300px, display_width_inches will be small.

        print(f"Calculated display width for DOCX: {display_width_inches:.2f} inches")

        if display_width_inches > 0.1 : # Ensure it's a sensible positive width
            # By only setting the width, Word will maintain the aspect ratio for height.
            doc.add_picture(image_path, width=Inches(display_width_inches))
            print(f"Added image {os.path.basename(image_path)} to DOCX with width {display_width_inches:.2f} inches. Height will be auto-scaled by Word.")
        else:
            print(f"Warning: Calculated display width for {os.path.basename(image_path)} is too small ({display_width_inches:.2f}). Skipping add_picture.")

    except FileNotFoundError:
        print(f"Error: Image file not found at {image_path} during DOCX processing.")
        return
    except Exception as e:
        print(f"Error processing image {image_path} for DOCX: {e}")
        return
    try:
        doc.save(docx_path)
        print(f"Word document saved successfully at: {docx_path}")
    except Exception as e:
        print(f"Error saving DOCX {docx_path}: {e}")

# --- Main Iteration Logic ---
for value in range(8, 30, 8):
    print(f"\n--- Processing for value: {value} ---")

    run_rhinocode_command(f"GetInteger {value} {value/4}")
    run_rhinocode_command("_SelAll")

    # Set display mode for cleaner render
    # IMPORTANT: Ensure "Rendered" mode in your Rhino has a SOLID WHITE background
    # Tools > Options > View > Display Modes > Rendered > Background > Solid color (White)
    run_rhinocode_command("_-SetDisplayMode _Mode=Rendered _Enter")

    stl_filename = f"model_{value}.stl"
    abs_stl_path = os.path.join(abs_output_dir, stl_filename)
    export_stl_command = f'_-Export "{abs_stl_path}" _EnterEnd'
    run_rhinocode_command(export_stl_command)

    run_rhinocode_command("_SelAll")
    run_rhinocode_command("_Zoom _Selected")
    run_rhinocode_command("_SelNone")

    png_filename = f"render_{value}.png"
    abs_png_path = os.path.join(abs_output_dir, png_filename)

    # ViewCaptureToFile: Still using the "chained input" style.
    # If options (Width, Height, TransparentBackground) are not applied,
    # a helper Rhino Python script is the most reliable alternative.
    view_capture_command = (
        f'_-ViewCaptureToFile '
        f'"{abs_png_path}" '
        f'Width=1920 '
        f'Height=1080 '
        f'Scale=1 '
        f'DrawGrid=No DrawAxes=No DrawWorldAxes=No DrawCPlaneAxes=No '
        f'TransparentBackground=Yes ' # This needs the display mode bg to be solid white
        f'_Enter'
    )
    run_rhinocode_command(view_capture_command)
    time.sleep(2)  # Give time for the PNG file to be written

    docx_filename = f"report_{value}.docx"
    abs_docx_path = os.path.join(abs_output_dir, docx_filename)

    # Add retry logic for PNG file
    max_retries = 3
    retry_count = 0
    while retry_count < max_retries:
        print(f"Checking for PNG at: {abs_png_path} (Attempt {retry_count + 1}/{max_retries})")
        if os.path.exists(abs_png_path):
            # Verify image size was as intended
            try:
                with Image.open(abs_png_path) as img_check:
                    print(f"PNG {os.path.basename(abs_png_path)} found. Dimensions: {img_check.size[0]}x{img_check.size[1]}px.")
                    break  # File exists and can be opened, proceed
            except Exception as e_img:
                print(f"Could not open or read PNG dimensions for {abs_png_path}: {e_img}")
        time.sleep(2)  # Wait before retry
        retry_count += 1

    if os.path.exists(abs_png_path):
        print(f"Proceeding with DOCX creation for {os.path.basename(abs_png_path)}.")
        create_docx_from_image(abs_png_path, abs_docx_path, value)
        time.sleep(1)  # Give time for DOCX creation
    else:
        print(f"PNG file {abs_png_path} NOT FOUND after {max_retries} attempts. Skipping DOCX creation.")
        print(f"Investigate Rhino logs for why '{png_filename}' was not saved or if options failed.")

    run_rhinocode_command('_SelAll _Delete _Enter')
    # Optional: Switch back to a default display mode if needed
    # run_rhinocode_command("_-SetDisplayMode _Mode=Shaded _Enter")


print(f"\n--- Iteration complete ({value}) ---")
print(f"All output files should be in: {abs_output_dir}")
